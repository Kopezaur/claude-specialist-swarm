"""
Run Card A (Deal Desk) end-to-end using the docx-creation skill.

This script plays both the coordinator and the four specialists locally
(no API calls), then synthesises their outputs into a single branded
Word document by following the skills/docx-creation/SKILL.md guide
exactly: Gryffindor palette, mandatory Harry Potter's Crew -> Hogwarts
header/footer/cover attribution, branded tables, and the three required
charts (approach, cost, timeline).

Output: outputs/hogwarts-proposal-v1.docx
"""

from __future__ import annotations

from pathlib import Path

import matplotlib

matplotlib.use("Agg")  # headless render
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Brand constants -- copied verbatim from skills/docx-creation/SKILL.md
# ---------------------------------------------------------------------------

FIRM_NAME = "Harry Potter's Crew"
CLIENT_NAME = "Hogwarts"

GRYFFINDOR_RED = RGBColor(0x7F, 0x09, 0x09)
GRYFFINDOR_GOLD = RGBColor(0xFF, 0xC5, 0x00)
GRYFFINDOR_DARK_RED = RGBColor(0x3C, 0x00, 0x00)
GRYFFINDOR_CREAM = RGBColor(0xFF, 0xF8, 0xE7)
BODY_BLACK = RGBColor(0x10, 0x10, 0x10)

CHART_PALETTE = ["#7F0909", "#FFC500", "#3C0000", "#B8860B", "#FFF8E7"]

OUTPUT_DIR = Path(__file__).parent / "outputs"
FIGURES_DIR = OUTPUT_DIR / "figures"
DOC_PATH = OUTPUT_DIR / "hogwarts-proposal-v2.docx"


# ---------------------------------------------------------------------------
# Skill helpers (verbatim from SKILL.md)
# ---------------------------------------------------------------------------

def _apply_base_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    body = normal.font
    body.name = "Calibri"
    body.size = Pt(11)
    body.color.rgb = BODY_BLACK

    # Tight paragraph spacing -- keeps the document compact.
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.15

    for level, size in [(0, 28), (1, 20), (2, 14), (3, 12)]:
        style_name = "Title" if level == 0 else f"Heading {level}"
        style = styles[style_name]
        h = style.font
        h.name = "Cambria"
        h.size = Pt(size)
        h.color.rgb = GRYFFINDOR_DARK_RED
        h.bold = True
        # Headings: a little breathing room above, almost none below.
        style.paragraph_format.space_before = Pt(8 if level else 0)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.keep_with_next = True


def _add_cover_page(doc: Document, title: str, subtitle: str) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(f"{FIRM_NAME}\nfor\n{CLIENT_NAME}")
    run.font.size = Pt(22)
    run.font.color.rgb = GRYFFINDOR_RED
    run.bold = True

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title)
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = GRYFFINDOR_DARK_RED
    title_run.bold = True

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(subtitle)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = GRYFFINDOR_GOLD
    sub_run.italic = True

    # Cover page stands alone; the rest of the document flows continuously.
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)



def _add_header(doc: Document) -> None:
    header_p = doc.sections[0].header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run(f"{FIRM_NAME}  ->  {CLIENT_NAME}")
    run.font.size = Pt(9)
    run.font.color.rgb = GRYFFINDOR_RED
    run.bold = True


def _add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = f" {field_code} "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.size = Pt(9)
    run.font.color.rgb = GRYFFINDOR_DARK_RED


def _add_footer(doc: Document) -> None:
    footer_p = doc.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    prefix = footer_p.add_run(
        f"{FIRM_NAME}  -  Prepared for {CLIENT_NAME}  -  Confidential  -  Page "
    )
    prefix.font.size = Pt(9)
    prefix.font.color.rgb = GRYFFINDOR_DARK_RED

    _add_field(footer_p, "PAGE")
    of_run = footer_p.add_run(" of ")
    of_run.font.size = Pt(9)
    of_run.font.color.rgb = GRYFFINDOR_DARK_RED
    _add_field(footer_p, "NUMPAGES")


def new_branded_doc(title: str, subtitle: str) -> Document:
    doc = Document()
    _apply_base_styles(doc)
    _add_cover_page(doc, title, subtitle)
    _add_header(doc)
    _add_footer(doc)
    return doc


def _shade_cell(cell, hex_fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def add_branded_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 2"

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = GRYFFINDOR_GOLD
        _shade_cell(cell, "7F0909")

    for r_idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = str(value)
            if r_idx % 2 == 0:
                _shade_cell(row_cells[c_idx], "FFF8E7")
    return table


def render_chart(fig, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    fig.savefig(
        out_path, dpi=200, bbox_inches="tight", facecolor="white", edgecolor="none"
    )
    plt.close(fig)
    return out_path


def embed_image(
    doc: Document, png_path: Path, caption: str, width_inches: float = 5.5
) -> None:
    doc.add_picture(str(png_path), width=Inches(width_inches))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(f"Figure - {caption}")
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = GRYFFINDOR_DARK_RED


# ---------------------------------------------------------------------------
# Required charts
# ---------------------------------------------------------------------------

def _styled_axes(ax):
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    for spine in ("left", "bottom"):
        ax.spines[spine].set_color("#7F0909")
    ax.tick_params(colors="#3C0000")
    ax.title.set_color("#3C0000")


def approach_chart(out_path: Path) -> Path:
    pillars = [
        "Discovery &\nArchitecture",
        "Lakehouse\nBuild",
        "Pensieve BI\nIntegration",
        "Predictive\nMaintenance ML",
        "Rollout &\nEnablement",
    ]
    effort = [10, 60, 25, 20, 30]
    fig, ax = plt.subplots(figsize=(7, 3.8))
    bars = ax.barh(
        pillars, effort, color=CHART_PALETTE[: len(pillars)],
        edgecolor="#3C0000", linewidth=1.2,
    )
    ax.invert_yaxis()
    ax.set_xlabel("Effort (working days)")
    ax.set_title("Proposed approach - five workstreams")
    for bar, days in zip(bars, effort):
        ax.text(bar.get_width() + 1, bar.get_y() + bar.get_height() / 2,
                f"{days}d", va="center", color="#3C0000", fontweight="bold")
    _styled_axes(ax)
    return render_chart(fig, out_path)


def cost_breakdown_chart(out_path: Path) -> Path:
    labels = ["Year 1\nplatform", "Year 2-3\nplatform", "Implementation\nservices", "Support &\nCSM"]
    values = [468, 936, 240, 180]  # in $K
    fig, ax = plt.subplots(figsize=(6.5, 3.8))
    bars = ax.bar(labels, values, color=CHART_PALETTE[: len(labels)],
                  edgecolor="#3C0000", linewidth=1.2)
    ax.set_ylabel("USD (thousands)")
    ax.set_title("3-year cost breakdown - illustrative")
    for bar, v in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 15,
                f"${v}K", ha="center", color="#3C0000", fontweight="bold")
    _styled_axes(ax)
    return render_chart(fig, out_path)


def timeline_chart(out_path: Path) -> Path:
    milestones = [
        ("Discovery & architecture", 0, 4),
        ("Lakehouse build", 4, 12),
        ("Pensieve BI integration", 8, 8),
        ("Predictive maintenance ML", 14, 6),
        ("Rollout & enablement", 18, 6),
        ("Steady state hand-over", 24, 2),
    ]
    fig, ax = plt.subplots(figsize=(7, 3.8))
    for i, (name, start, duration) in enumerate(milestones):
        ax.barh(i, duration, left=start,
                color=CHART_PALETTE[i % len(CHART_PALETTE)],
                edgecolor="#3C0000", linewidth=1.2)
        ax.text(start + duration + 0.3, i, f"w{start}-w{start + duration}",
                va="center", color="#3C0000", fontsize=8)
    ax.set_yticks(range(len(milestones)))
    ax.set_yticklabels([m[0] for m in milestones])
    ax.invert_yaxis()
    ax.set_xlabel("Programme weeks")
    ax.set_title("Implementation timeline - 26 weeks")
    _styled_axes(ax)
    return render_chart(fig, out_path)


# ---------------------------------------------------------------------------
# Specialist outputs (the coordinator synthesises these in real life;
# here they are baked-in summaries the local "specialists" produced).
# ---------------------------------------------------------------------------

PRICING_RECOMMENDATION = (
    "Enterprise tier at $720K list. Hogwarts' workload (real-time ingest from 40,000 "
    "magical artefacts, ~280TB and growing) sits comfortably above the $500K threshold "
    "where Enterprise is the only sensible tier. We recommend a 3-year initial term, "
    "annual upfront billing on Net 30, and a 25% strategic discount band (multi-year + "
    "marquee logo). Net annual fee: $540K. Pilot fees credited to Year 1. We will NOT "
    "accept the RFP's MFN clause, Net 90 payment terms, or uncapped liability - those "
    "are non-negotiable per the pricing playbook."
)

LEGAL_FLAGS = [
    ("Liability cap", "BLOCKER",
     "RFP demands uncapped liability for data breach. Our cyber policy caps at 24 months of fees - uncapped voids coverage. "
     "Counter: 24 months of fees, with mutual carve-outs for IP infringement and gross negligence."),
    ("Termination for convenience", "BLOCKER",
     "30-day termination for any reason, no penalty. Counter: 90 days' notice after the first 12 months of the initial term."),
    ("Audit rights", "BLOCKER",
     "Up to 4 audits/year without notice, vendor-paid. Counter: 1 audit/year, 30 days' notice, audit confidentiality required, Hogwarts pays beyond the first."),
    ("Service levels", "BLOCKER",
     "99.99% uptime with immediate termination on any miss. Counter: 99.95% Enterprise SLA, service credits up to 30% of monthly fees as sole remedy."),
    ("Most Favoured Nation", "BLOCKER",
     "MFN pricing for the life of the contract. Counter: remove entirely - the single most poisonous clause for SaaS."),
    ("IP assignment", "NEGOTIABLE",
     "RFP demands assignment of all work product. Counter: Hogwarts owns customer-specific configurations; HP Crew retains IP in the underlying service and any general-purpose enhancements."),
    ("Payment terms", "NEGOTIABLE",
     "Net 90 demanded. Counter: Net 30 default, Net 60 acceptable given Hogwarts' standing and 3-year commitment."),
    ("Subprocessors", "NEGOTIABLE",
     "Pre-approval of every subprocessor. Counter: published subprocessor list with 30 days' notice of additions and a right to object."),
]

TECHNICAL_FIT = (
    "Our Enterprise Data Platform covers Hogwarts' core requirements: lakehouse "
    "architecture on open formats (Parquet + Delta), Azure-native deployment with EU "
    "primary and US East secondary regions, real-time ingest at 80,000 events/sec "
    "(headroom 2x peak), and native Pensieve BI (Power BI) integration for the 600 "
    "analyst seat estate. Predictive maintenance ML pipelines map to our MLflow-backed "
    "feature store. The only gap: data residency for EU customer data is satisfied by "
    "default; we will need to confirm Munich R&D's specific dataset classification "
    "during discovery."
)

COMPETITIVE_READ = (
    "Hogwarts named Databricks, Snowflake, and Microsoft Fabric. Strongest likely "
    "competitor: Microsoft Fabric (Hogwarts is an Azure + Pensieve BI shop, E5 already "
    "deployed). Our best two angles: (1) honest TCO including Microsoft consulting "
    "hours, (2) maturity - HP Crew has 8 years on this; Fabric has 18 months. Trap to "
    "avoid: do not compete on Pensieve BI integration - they own it, we integrate. "
    "Against Databricks, lead with TCO and analyst time-to-insight. Against Snowflake, "
    "lead with workload coverage (real-time, unstructured) and ML-native architecture."
)


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------

def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FIGURES_DIR.mkdir(parents=True, exist_ok=True)

    doc = new_branded_doc(
        title="Modernising the Hogwarts Records & Magical Artefact Platform",
        subtitle="A proposal from Harry Potter's Crew - May 2026",
    )

    # 2. Table of contents
    doc.add_heading("Table of Contents", level=1)
    toc_p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-2" \h \z \u')
    placeholder = OxmlElement("w:r")
    placeholder_text = OxmlElement("w:t")
    placeholder_text.text = "Right-click and select 'Update Field' in Word to populate this table of contents."
    placeholder.append(placeholder_text)
    fld.append(placeholder)
    toc_p._p.append(fld)

    # 3. Executive Summary
    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "Hogwarts' RFP calls for an enterprise data platform to replace the current "
        "patchwork of on-premises archives and ad-hoc cloud analytics. The platform "
        "must support real-time ingest from 40,000 magical artefacts in the field, "
        "batch ETL from 30+ internal sources, BI for 600 analysts, self-service "
        "preparation for 150 data engineers, and a planned predictive-maintenance "
        "ML programme - all on an Azure-primary, EU-resident footprint with native "
        "Pensieve BI (Power BI) integration."
    )
    doc.add_paragraph(
        "Harry Potter's Crew proposes our Enterprise tier on a 3-year initial term "
        "with a 2-year renewal option, delivered in a 26-week implementation programme "
        "across five workstreams. We have addressed every functional requirement, "
        "flagged the eight contractual positions that must move before signature, and "
        "delivered a 3-year total cost of $1.82M against a market-clearing benchmark."
    )
    why_p = doc.add_paragraph()
    why_p.add_run("Why Harry Potter's Crew:").bold = True
    doc.add_paragraph(
        "Eight years operating lakehouse architectures at the scale Hogwarts requires.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "ML-native platform - predictive maintenance is a first-class workload, not a bolt-on.",
        style="List Bullet",
    )
    doc.add_paragraph(
        "Honest, fixed TCO with no escalators - the only bid you can trust at year 3.",
        style="List Bullet",
    )

    # 4. Understanding of Hogwarts' needs
    doc.add_heading("2. Understanding of Hogwarts' Needs", level=1)
    doc.add_paragraph(
        "We have read the RFP closely. The platform must consolidate the legacy "
        "parchment-and-quill archive (Teradata-equivalent), serve 600 analysts on "
        "Pensieve BI, and provide a foundation for the planned predictive-maintenance "
        "programme on the 40,000-artefact estate. Three constraints stand out:"
    )
    add_branded_table(
        doc,
        headers=["Requirement", "Hogwarts' position", "Our coverage"],
        rows=[
            ["Native Pensieve BI integration", "Non-negotiable - 600 active users", "Full coverage; certified Pensieve BI connector"],
            ["Multi-region deployment", "EU primary, US East secondary", "Azure-native; EU data residency by default"],
            ["Scale headroom", "280TB today, 12TB/month growth, 80K events/sec peak", "Tested to 160K events/sec; 2x headroom on day 1"],
            ["Lakehouse + open formats", "Parquet, Delta, Iceberg for portability", "Delta-first; Iceberg read/write GA"],
            ["Predictive maintenance ML", "Planned, not active", "MLflow-backed feature store ready when programme starts"],
        ],
    )

    # 5. Proposed approach
    doc.add_heading("3. Proposed Approach", level=1)
    doc.add_paragraph(
        "We propose a 26-week implementation programme across five workstreams, "
        "running partly in parallel. Discovery anchors the architecture; the lakehouse "
        "build and Pensieve BI integration run concurrently from week 4; predictive "
        "maintenance ML begins once the feature store has 8 weeks of clean data."
    )
    embed_image(
        doc,
        approach_chart(FIGURES_DIR / "approach.png"),
        caption="Workstream effort - five pillars of the engagement",
    )
    doc.add_paragraph(
        "Each workstream is led by a named HP Crew partner. Hogwarts will have a "
        "dedicated Customer Success Manager from day 1 and weekly steering with the "
        "Chief Data Officer's office."
    )

    # 6. Commercials
    doc.add_heading("4. Commercials", level=1)
    doc.add_paragraph(
        "Pricing is built on our Enterprise tier at $720K annual list. Hogwarts "
        "qualifies for the Strategic 25% discount band (3-year commitment, marquee "
        "logo, reference customer agreement). Net annual platform fee: $540K. "
        "Implementation services and 24/7 support are quoted separately."
    )
    add_branded_table(
        doc,
        headers=["Workstream", "Term", "Fee (USD)"],
        rows=[
            ["Enterprise platform - Year 1", "12 months", "$468,000"],
            ["Enterprise platform - Year 2 & 3", "24 months", "$936,000"],
            ["Implementation services (26 weeks)", "One-time", "$240,000"],
            ["24/7 support + dedicated CSM", "36 months", "$180,000"],
            ["TOTAL - 3 year programme", "-", "$1,824,000"],
        ],
    )
    embed_image(
        doc,
        cost_breakdown_chart(FIGURES_DIR / "cost-breakdown.png"),
        caption="3-year cost breakdown by category",
    )
    doc.add_paragraph(
        "Pilot/POC fees are credited toward Year 1. We offer a 30-day acceptance "
        "testing window. We will not accept the RFP's MFN clause, Net 90 payment "
        "terms, or uncapped breach liability - these are non-negotiable and our "
        "Counter-Positions are detailed in Section 8."
    )

    # 7. Timeline
    doc.add_heading("5. Timeline & Milestones", level=1)
    doc.add_paragraph(
        "The 26-week programme is sequenced to put a working lakehouse in front of "
        "Hogwarts' analysts by week 12 and to begin the predictive-maintenance ML "
        "programme by week 14 once the feature store has stabilised."
    )
    embed_image(
        doc,
        timeline_chart(FIGURES_DIR / "timeline.png"),
        caption="Implementation timeline - 26 working weeks",
    )

    # 8. Why Harry Potter's Crew
    doc.add_heading("6. Why Harry Potter's Crew", level=1)
    doc.add_paragraph(
        "Hogwarts named Databricks, Snowflake, and Microsoft Fabric in the RFP. Each "
        "is a credible competitor. Our differentiation is not headline price - it is "
        "predictable total cost of ownership over the full 5-year horizon, maturity at "
        "Hogwarts' scale, and an ML-native architecture that does not treat predictive "
        "maintenance as a bolt-on."
    )
    add_branded_table(
        doc,
        headers=["Competitor", "Their strength", "Our angle"],
        rows=[
            ["Microsoft Fabric", "Bundled in E5, owns Pensieve BI", "Honest TCO including Microsoft consulting; 8 years vs. 18 months of maturity"],
            ["Databricks", "Lakehouse + ML breadth", "TCO predictability; analyst time-to-insight"],
            ["Snowflake", "Analyst experience, data sharing", "Workload coverage (real-time, unstructured); ML-native"],
            ["Regional unnamed vendor", "Unknown - likely price-led", "Vendor financial stability; reference customers at Hogwarts' scale"],
        ],
    )

    # 9. Legal counter-positions (was originally appendix; promoted because it's significant)
    doc.add_heading("7. Contractual Counter-Positions", level=1)
    doc.add_paragraph(
        "Our Legal Reviewer flagged eight RFP clauses requiring movement before "
        "signature. Five are BLOCKERS (cannot be insured around or commercially "
        "justified) and three are NEGOTIABLE."
    )
    add_branded_table(
        doc,
        headers=["Clause", "Severity", "HP Crew counter-position"],
        rows=[[name, sev, counter] for name, sev, counter in LEGAL_FLAGS],
    )

    # 10. Appendix
    doc.add_heading("8. Appendix - Assumptions, Risks & Glossary", level=1)
    doc.add_heading("8.1 Assumptions baked into this proposal", level=2)
    for note in [
        "Currency is USD; conversion to galleons at the prevailing Gringotts rate is the client's responsibility.",
        "All headcount and scale figures (40,000 artefacts, 280TB, 600 analysts) are taken from the RFP at face value.",
        "The Capability Matrix attachment is assumed to be returned separately as the RFP instructs.",
        "Pricing is illustrative for this synthetic exercise and would be confirmed by the Deal Desk before issue.",
        "Pensieve BI is used as the Hogwarts-themed name for Power BI throughout this document.",
    ]:
        doc.add_paragraph(note, style="List Bullet")

    doc.add_heading("8.2 Top three risks", level=2)
    add_branded_table(
        doc,
        headers=["Risk", "Likelihood", "Mitigation"],
        rows=[
            ["Legal blockers (esp. uncapped liability, MFN) not movable", "Medium", "Escalate to Hogwarts general counsel in Week 1"],
            ["Pensieve BI semantic-model migration scope grows", "Medium", "Bound the scope to top-100 reports; tail-end remains client-led"],
            ["Predictive maintenance feature store delayed by data quality", "Low-Medium", "Treat Weeks 4-8 as a quality gate; do not start ML until clean"],
        ],
    )

    doc.save(DOC_PATH)
    return DOC_PATH


# ---------------------------------------------------------------------------
# Post-save branding checks (as the skill demands)
# ---------------------------------------------------------------------------

def verify_branding(path: Path) -> dict:
    doc = Document(path)
    footer_text = "\n".join(
        p.text for s in doc.sections for p in s.footer.paragraphs
    )
    header_text = "\n".join(
        p.text for s in doc.sections for p in s.header.paragraphs
    )
    image_count = sum(
        1
        for rel in doc.part._rels.values()
        if "image" in rel.target_ref
    )
    table_count = len(doc.tables)
    heading_count = sum(
        1 for p in doc.paragraphs if (p.style.name or "").startswith("Heading")
    )
    return {
        "firm_in_footer": FIRM_NAME in footer_text,
        "client_in_footer": CLIENT_NAME in footer_text,
        "firm_in_header": FIRM_NAME in header_text,
        "client_in_header": CLIENT_NAME in header_text,
        "image_count": image_count,
        "table_count": table_count,
        "heading_count": heading_count,
    }


def main() -> None:
    path = build_document()
    checks = verify_branding(path)
    print(f"Saved: {path.resolve()}")
    print("Branding checks:")
    for k, v in checks.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
